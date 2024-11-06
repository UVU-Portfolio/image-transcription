import time
import logging
import queue
import threading
import shutil
from pathlib import Path
from typing import Optional, List
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import subprocess
from docx import Document
from datetime import datetime
from tqdm import tqdm

# Set up root logger to output to both file and console
def setup_root_logger():
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)
    
    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    console_handler.setFormatter(console_formatter)
    logger.addHandler(console_handler)
    
    return logger

logger = setup_root_logger()

@dataclass
class Config:
    """Configuration settings for the OCR processor"""
    base_dir: Path = Path('/app/input')
    monitored_dir: Path = Path('/app/input/images')
    output_dir: Path = Path('/app/output')
    pix2tex_path: Path = Path('/app/pix2tex/run.py')
    max_workers: int = 4
    supported_formats: tuple = ('.png', '.jpg', '.jpeg')
    batch_size: int = 10
    test_mode: bool = True

    def __post_init__(self):
        logger.info(f"Initialized config with:")
        logger.info(f"  Base directory: {self.base_dir}")
        logger.info(f"  Monitored directory: {self.monitored_dir}")
        logger.info(f"  Output directory: {self.output_dir}")
        logger.info(f"  Pix2tex path: {self.pix2tex_path}")
        logger.info(f"  Supported formats: {self.supported_formats}")

    def setup_directories(self):
        """Create necessary directories"""
        logger.info("Creating necessary directories")
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self.monitored_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Created directories: {self.base_dir}, {self.monitored_dir}, {self.output_dir}")

def is_image_file(file_path: Path) -> bool:
    """Check if a file is a supported image file"""
    return file_path.is_file() and file_path.suffix.lower() in Config.supported_formats

class TranscriptionQueue:
    def __init__(self, batch_size: int):
        self.queue = queue.Queue()
        self.batch_size = batch_size
        self.lock = threading.Lock()
        logger.info(f"Initialized TranscriptionQueue with batch size: {batch_size}")
    
    def add_task(self, image_path: Path):
        logger.info(f"Adding task to queue: {image_path}")
        self.queue.put(image_path)
    
    def get_batch(self) -> List[Path]:
        batch = []
        try:
            while len(batch) < self.batch_size:
                item = self.queue.get_nowait()
                batch.append(item)
        except queue.Empty:
            if batch:
                logger.debug(f"Got partial batch of size {len(batch)}")
        return batch

class OCRProcessor:
    def __init__(self, config: Config):
        self.config = config
        self.logger = self._setup_logger()
        self.queue = TranscriptionQueue(config.batch_size)
        self.executor = ThreadPoolExecutor(max_workers=config.max_workers)
        logger.info("Starting processing thread...")
        self.processing_thread = threading.Thread(target=self._process_queue, daemon=True)
        self.processing_thread.start()
        logger.info("Processing thread started")

    def _setup_logger(self) -> logging.Logger:
        logger = logging.getLogger('OCRProcessor')
        logger.setLevel(logging.DEBUG)
        
        # File handler
        file_handler = logging.FileHandler(self.config.output_dir / 'ocr_processor.log')
        file_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(file_formatter)
        logger.addHandler(file_handler)
        
        return logger

    def validate_image(self, image_path: Path) -> bool:
        """Validate that the file is a proper image file"""
        if not image_path.is_file():
            logger.error(f"Not a file: {image_path}")
            return False
        if not image_path.suffix.lower() in self.config.supported_formats:
            logger.error(f"Unsupported format: {image_path}")
            return False
        try:
            # Import PIL only when needed
            from PIL import Image
            img = Image.open(image_path)
            img.verify()
            return True
        except Exception as e:
            logger.error(f"Invalid image file {image_path}: {e}")
            return False

    def transcribe_image(self, image_path: Path) -> Optional[str]:
        try:
            logger.info(f"Starting transcription of image: {image_path}")
            
            if not self.validate_image(image_path):
                return None
            
            logger.info(f"Image validated successfully: {image_path}")
            command = f'python3 {self.config.pix2tex_path} --image {image_path}'
            logger.info(f"Running command: {command}")
            
            result = subprocess.run(
                command,
                shell=True,
                capture_output=True,
                text=True,
                check=True,
                timeout=30
            )
            
            if result.stderr:
                logger.warning(f"Command stderr: {result.stderr}")
            
            output = result.stdout.strip()
            if not output:
                logger.warning("Command produced no output")
                return None
                
            logger.info(f"Transcription completed for {image_path}")
            logger.debug(f"Transcription output: {output}")
            return output
            
        except subprocess.CalledProcessError as e:
            logger.error(f"Transcription error for {image_path}: {e}")
            logger.error(f"Command stdout: {e.stdout}")
            logger.error(f"Command stderr: {e.stderr}")
            return None
        except subprocess.TimeoutExpired:
            logger.error(f"Transcription timeout for {image_path}")
            return None
        except Exception as e:
            logger.error(f"Unexpected error transcribing {image_path}: {e}", exc_info=True)
            return None

    def _process_queue(self):
        logger.info("Queue processing started")
        while True:
            batch = self.queue.get_batch()
            if batch:
                logger.info(f"Processing batch of {len(batch)} images")
                self._process_batch(batch)
            time.sleep(1)

    def _process_batch(self, batch: List[Path]):
        doc_path = self._get_batch_doc_path()
        logger.info(f"Creating new batch document: {doc_path}")
        doc = Document()
        doc.add_heading('Batch OCR Transcriptions', level=0)

        with tqdm(total=len(batch), desc="Processing images") as pbar:
            for image_path in batch:
                logger.info(f"Processing image from batch: {image_path}")
                transcription = self.transcribe_image(image_path)
                if transcription:
                    self._append_to_doc(doc, image_path, transcription)
                    logger.info(f"Successfully transcribed: {image_path}")
                else:
                    logger.warning(f"Failed to transcribe: {image_path}")
                pbar.update(1)

        try:
            logger.info(f"Saving batch document to: {doc_path}")
            doc.save(doc_path)
            logger.info(f"Successfully saved batch document: {doc_path}")
        except Exception as e:
            logger.error(f"Error saving document: {e}", exc_info=True)

    def _get_batch_doc_path(self) -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return self.config.output_dir / f'batch_transcriptions_{timestamp}.docx'

    def _append_to_doc(self, doc: Document, image_path: Path, transcription: str):
        try:
            logger.debug(f"Appending transcription for {image_path}")
            doc.add_heading(f'Transcription for {image_path.name}', level=1)
            doc.add_paragraph(transcription)
        except Exception as e:
            logger.error(f"Error appending to document: {e}", exc_info=True)

    def process_existing_files(self):
        """Process any existing files in the monitored directory"""
        logger.info(f"Checking for existing files in {self.config.monitored_dir}")
        if not self.config.monitored_dir.exists():
            logger.error(f"Monitored directory does not exist: {self.config.monitored_dir}")
            return

        existing_files = [
            f for f in self.config.monitored_dir.glob('*')
            if is_image_file(f)
        ]
        
        logger.info(f"Found {len(existing_files)} existing files to process")
        for file_path in existing_files:
            logger.info(f"Adding existing file to queue: {file_path}")
            self.queue.add_task(file_path)

    def run_test(self) -> bool:
        """Run a test transcription"""
        logger.info("Running test transcription")
        
        # Find image files
        test_files = [
            f for f in self.config.monitored_dir.glob('*')
            if is_image_file(f)
        ]
        
        logger.info(f"Found {len(test_files)} image files in monitored directory")
        
        if not test_files:
            logger.warning("No test image files found - skipping test")
            return True  # Changed to return True when no files found

        test_file = test_files[0]
        logger.info(f"Selected test file: {test_file}")
        
        transcription = self.transcribe_image(test_file)
        if not transcription:
            logger.error("Test transcription failed")
            return False

        test_doc_path = self.config.output_dir / 'test_transcription.docx'
        try:
            logger.info(f"Saving test transcription to: {test_doc_path}")
            doc = Document()
            doc.add_paragraph(transcription)
            doc.save(test_doc_path)
            logger.info(f"Test successful, saved to: {test_doc_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving test document: {e}", exc_info=True)
            return False

class ImageEventHandler(FileSystemEventHandler):
    def __init__(self, processor: OCRProcessor, config: Config):
        self.processor = processor
        self.config = config
        logger.info("Initialized ImageEventHandler")

    def on_created(self, event):
        if event.is_directory:
            return
        
        file_path = Path(event.src_path)
        logger.info(f"New file detected: {file_path}")
        
        if is_image_file(file_path):
            logger.info(f"Processing new image file: {file_path}")
            self.processor.queue.add_task(file_path)
        else:
            logger.warning(f"Ignoring unsupported file: {file_path}")

def main():
    logger.info("Starting OCR monitor")
    config = Config()
    
    try:
        # Set up directories
        logger.info("Setting up directories")
        config.setup_directories()
        
        # Move any images from the root directory to the images subdirectory
        logger.info("Moving any images from root directory to images subdirectory")
        for file_path in config.base_dir.glob('*'):
            if is_image_file(file_path) and file_path.parent == config.base_dir:
                target_path = config.monitored_dir / file_path.name
                logger.info(f"Moving {file_path} to {target_path}")
                shutil.move(str(file_path), str(target_path))
        
        logger.info("Initializing OCR processor")
        processor = OCRProcessor(config)
        
        if config.test_mode:
            logger.info("Running in test mode")
            if not processor.run_test():
                logger.error("Test failed, stopping execution")
                raise RuntimeError("Test failed, stopping execution")

        logger.info("Processing existing files")
        processor.process_existing_files()

        logger.info("Setting up file system observer")
        event_handler = ImageEventHandler(processor, config)
        observer = Observer()
        observer.schedule(event_handler, path=str(config.monitored_dir), recursive=False)
        observer.start()

        print(f"Monitoring directory: {config.monitored_dir}")
        print(f"Output directory: {config.output_dir}")
        
        logger.info("Monitor is now running")
        while True:
            time.sleep(1)

    except KeyboardInterrupt:
        logger.info("Received keyboard interrupt, shutting down")
        observer.stop()
        observer.join()
    except Exception as e:
        logger.error(f"Fatal error: {e}", exc_info=True)
        raise

if __name__ == "__main__":
    main()